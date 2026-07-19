from pathlib import Path

from docx import Document

from docflow_docx.writer import write_docx_from_html


def test_write_docx_preserves_wrapped_paragraphs(tmp_path: Path) -> None:
    html = (
        '<div class="docx-document" style="font-size:11pt">'
        "<p>Абзац один</p><p>Абзац два</p>"
        "</div>"
    )
    path = tmp_path / "wrapped.docx"
    write_docx_from_html(path, html)

    doc = Document(path)
    texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

    assert texts == ["Абзац один", "Абзац два"]


def test_write_docx_preserves_canvas_wrapper(tmp_path: Path) -> None:
    html = (
        '<div class="docx-canvas">'
        '<div class="docx-document docx-editable">'
        "<p>Перший</p><p>Другий</p>"
        "</div></div>"
    )
    path = tmp_path / "canvas.docx"
    write_docx_from_html(path, html)

    doc = Document(path)
    texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

    assert texts == ["Перший", "Другий"]
