import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


def write_docx_from_html(path: Path, html_content: str) -> None:
    write_docx_from_pages(path, [{"html": html_content}])


def write_docx_from_pages(path: Path, pages: list[dict]) -> None:
    doc = Document()

    for index, page in enumerate(pages):
        if index > 0:
            doc.add_page_break()

        html_content = page.get("html", "").strip()
        if not html_content:
            doc.add_paragraph("")
            continue

        soup = BeautifulSoup(f"<div>{html_content}</div>", "html.parser")
        wrapper = soup.find("div")
        if wrapper:
            for child in wrapper.children:
                _append_node(doc, child)

    doc.save(path)


def _append_node(doc: Document, node) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text.strip():
            doc.add_paragraph(text.strip())
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower()

    if name == "div":
        classes = node.get("class") or []
        if "docx-page-break" in classes:
            paragraph = doc.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            return

        for child in node.children:
            _append_node(doc, child)
        return

    if name in {"p", "h1", "h2", "h3", "h4", "h5", "h6"}:
        _add_paragraph_from_tag(doc, node, name)
        return

    if name in {"ul", "ol"}:
        for li in node.find_all("li", recursive=False):
            _add_paragraph_from_tag(doc, li, "li")
        return

    if name == "table":
        _add_table(doc, node)
        return

    if name == "br":
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_break()
        return

    if name == "img":
        return

    child_tags = [child for child in node.children if isinstance(child, Tag)]
    if child_tags:
        for child in node.children:
            _append_node(doc, child)
        return

    if node.get_text(strip=True):
        _add_paragraph_from_tag(doc, node, "p")


def _add_paragraph_from_tag(doc: Document, node: Tag, tag_name: str) -> None:
    paragraph = doc.add_paragraph()
    if tag_name.startswith("h") and len(tag_name) == 2 and tag_name[1].isdigit():
        paragraph.style = f"Heading {tag_name[1]}"

    style = _style_dict(node.get("style", ""))
    _apply_paragraph_style(paragraph, style)

    if "line-height" in style:
        pf = paragraph.paragraph_format
        try:
            pf.line_spacing = float(style["line-height"])
        except ValueError:
            pass

    _append_runs(paragraph, node)


def _append_runs(paragraph, node: Tag, inherited: dict | None = None) -> None:
    inherited = inherited or {}

    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = paragraph.add_run(text)
                _apply_run_style(run, inherited)
            continue

        if not isinstance(child, Tag):
            continue

        name = child.name.lower()
        style = {**inherited, **_style_dict(child.get("style", ""))}

        if name in {"strong", "b"}:
            style["bold"] = True
            _append_runs(paragraph, child, style)
        elif name in {"em", "i"}:
            style["italic"] = True
            _append_runs(paragraph, child, style)
        elif name == "u":
            style["underline"] = True
            _append_runs(paragraph, child, style)
        elif name in {"s", "strike"}:
            style["text-decoration"] = "line-through"
            _append_runs(paragraph, child, style)
        elif name == "sup":
            style["superscript"] = True
            _append_runs(paragraph, child, style)
        elif name == "sub":
            style["subscript"] = True
            _append_runs(paragraph, child, style)
        elif name == "span":
            _append_runs(paragraph, child, style)
        elif name == "br":
            paragraph.add_run().add_break()
        elif name == "img":
            continue
        else:
            text = child.get_text()
            if text:
                run = paragraph.add_run(text)
                _apply_run_style(run, style)


def _add_table(doc: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr", recursive=False)
    if not rows:
        rows = table_tag.find_all("tr")
    if not rows:
        return

    cols_count = max(len(row.find_all(["td", "th"], recursive=False)) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols_count)
    table.style = "Table Grid"

    for row_index, row_tag in enumerate(rows):
        cells = row_tag.find_all(["td", "th"], recursive=False)
        for col_index in range(cols_count):
            cell = table.rows[row_index].cells[col_index]
            cell.text = ""
            if col_index < len(cells):
                _fill_cell(cell, cells[col_index])


def _fill_cell(cell, cell_tag: Tag) -> None:
    blocks = cell_tag.find_all(["p", "h1", "h2", "h3", "h4", "h5", "h6", "li"], recursive=False)

    if blocks:
        for index, block in enumerate(blocks):
            paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            tag_name = block.name.lower()
            if tag_name.startswith("h") and len(tag_name) == 2 and tag_name[1].isdigit():
                paragraph.style = f"Heading {tag_name[1]}"
            style = _style_dict(block.get("style", ""))
            _apply_paragraph_style(paragraph, style)
            _append_runs(paragraph, block)
        return

    text = cell_tag.get_text().replace("\xa0", " ").strip()
    if text:
        _append_runs(cell.paragraphs[0], cell_tag)


def _apply_paragraph_style(paragraph, style: dict) -> None:
    pf = paragraph.paragraph_format
    align = style.get("text-align")
    if align == "center":
        pf.alignment = 1
    elif align == "right":
        pf.alignment = 2
    elif align == "justify":
        pf.alignment = 3

    for css_key, attr in (
        ("margin-left", "left_indent"),
        ("margin-right", "right_indent"),
        ("text-indent", "first_line_indent"),
        ("margin-top", "space_before"),
        ("margin-bottom", "space_after"),
    ):
        if css_key in style:
            setattr(pf, attr, Pt(_parse_pt(style[css_key])))


def _apply_run_style(run, style: dict) -> None:
    font = run.font
    if style.get("bold"):
        font.bold = True
    if style.get("italic"):
        font.italic = True
    if style.get("underline") or ("text-decoration" in style and "underline" in style["text-decoration"]):
        font.underline = True
    if "text-decoration" in style and "line-through" in style["text-decoration"]:
        font.strike = True
    if style.get("superscript"):
        font.superscript = True
    if style.get("subscript"):
        font.subscript = True
    if "font-size" in style:
        font.size = Pt(_parse_pt(style["font-size"]))
    if "font-family" in style:
        font.name = style["font-family"].strip("'\"")
    if "color" in style:
        color = style["color"].lstrip("#")
        if len(color) == 6:
            font.color.rgb = RGBColor.from_string(color)


def _style_dict(style: str) -> dict:
    result: dict = {}
    for chunk in style.split(";"):
        if ":" not in chunk:
            continue
        key, value = chunk.split(":", 1)
        result[key.strip()] = value.strip()
    return result


def _parse_pt(value: str) -> float:
    match = re.search(r"([\d.]+)", value)
    return float(match.group(1)) if match else 11.0
